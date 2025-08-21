import sys
import os
import re
import tempfile
import subprocess
import shutil
from docx import Document
from docx.shared import Inches
from PyQt6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QPushButton, QTextEdit, QFileDialog,
                             QLabel, QHBoxLayout, QLineEdit, QMessageBox, QProgressBar,
                             QScrollArea, QDialog, QGridLayout, QTabWidget)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QPixmap, QClipboard
from google import genai
from PIL import Image
import io

class ImagePreviewDialog(QDialog):
    def __init__(self, images, parent=None):
        super().__init__(parent)
        self.images = images
        self.initUI()

    def initUI(self):
        self.setWindowTitle(f'Xem trước {len(self.images)} hình ảnh')
        self.setGeometry(200, 200, 800, 600)

        layout = QVBoxLayout()

        # Create scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)

        # Create widget to hold images
        content_widget = QWidget()
        grid_layout = QGridLayout(content_widget)

        # Add images to grid (2 columns)
        row = 0
        col = 0
        for i, img_info in enumerate(self.images):
            try:
                # Create label for image
                img_label = QLabel()
                pixmap = QPixmap(img_info['path'])

                # Scale image to fit
                scaled_pixmap = pixmap.scaled(300, 300, Qt.AspectRatioMode.KeepAspectRatio,
                                            Qt.TransformationMode.SmoothTransformation)
                img_label.setPixmap(scaled_pixmap)
                img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

                # Create info label
                info_text = f"HÌNH ẢNH {i+1}\nKích thước: {img_info.get('size_info', 'N/A')}\nNguồn: {img_info.get('source', 'N/A')}"
                info_label = QLabel(info_text)
                info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                info_label.setStyleSheet("font-weight: bold; margin: 5px; padding: 5px; border: 1px solid #ccc; border-radius: 3px;")

                # Add to grid
                img_container = QVBoxLayout()
                img_container.addWidget(info_label)
                img_container.addWidget(img_label)

                container_widget = QWidget()
                container_widget.setLayout(img_container)

                grid_layout.addWidget(container_widget, row, col)

                # Move to next position
                col += 1
                if col >= 2:  # 2 columns
                    col = 0
                    row += 1

            except Exception as e:
                print(f"Error loading image {i+1}: {e}")

        scroll.setWidget(content_widget)
        layout.addWidget(scroll)

        # Close button
        close_button = QPushButton('Đóng')
        close_button.clicked.connect(self.accept)
        layout.addWidget(close_button)

        self.setLayout(layout)

class ConversionThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, client, uploaded_file, prompt):
        super().__init__()
        self.client = client
        self.uploaded_file = uploaded_file
        self.prompt = prompt
        self.max_retries = 3
        self.retry_delay = 60
        self.is_running = True

    def run(self):
        for attempt in range(self.max_retries):
            if not self.is_running:
                return
            try:
                # Simulate progress
                for i in range(0, 51):
                    if not self.is_running:
                        return
                    self.progress.emit(i)
                    self.msleep(50)

                # Progress simulation during generation
                for i in range(51, 101):
                    if not self.is_running:
                        return
                    self.progress.emit(i)
                    self.msleep(30)

                # Generate content using new simplified API
                response = self.client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[self.uploaded_file, self.prompt],
                )

                self.finished.emit(response.text)
                return

            except Exception as e:
                if "429" in str(e) and attempt < self.max_retries - 1:
                    self.error.emit(f"Rate limit exceeded. Retrying in {self.retry_delay} seconds...")
                    self.msleep(self.retry_delay * 1000)
                else:
                    self.error.emit(str(e))
                    return

    def stop(self):
        self.is_running = False

class WordTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_converter = parent
        self.initUI()

    def initUI(self):
        layout = QVBoxLayout()

        # Title
        title_label = QLabel("PDF to Word Converter")
        title_label.setStyleSheet("font-size: 16px; font-weight: bold; margin: 10px;")
        layout.addWidget(title_label)

        # Convert button
        self.convert_button = QPushButton('Convert PDF to Text')
        self.convert_button.clicked.connect(self.convert_pdf_to_text)
        self.convert_button.setEnabled(False)
        layout.addWidget(self.convert_button)

        # Export buttons
        export_layout = QHBoxLayout()
        self.export_word_button = QPushButton('Export to Word Document (python-docx)')
        self.export_word_button.clicked.connect(self.export_to_word)
        self.export_word_button.setEnabled(False)

        self.export_pandoc_button = QPushButton('Export to Word with Pandoc (Math formulas)')
        self.export_pandoc_button.clicked.connect(self.export_to_word_pandoc)
        self.export_pandoc_button.setEnabled(False)

        export_layout.addWidget(self.export_word_button)
        export_layout.addWidget(self.export_pandoc_button)
        layout.addLayout(export_layout)

        # Result display
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        layout.addWidget(QLabel('Results:'))
        layout.addWidget(self.result_text)

        self.setLayout(layout)

    def convert_pdf_to_text(self):
        if not self.parent_converter.client:
            QMessageBox.warning(self, "Error", "Please set the API Key first.")
            return

        if not self.parent_converter.uploaded_file:
            self.result_text.setText("Please upload a PDF file first.")
            return

        prompt = """
        Hãy nhận diện và gõ lại [CHÍNH XÁC] toàn bộ nội dung PDF thành văn bản, bao gồm tất cả công thức Toán học được bọc trong dấu $.

        [YÊU CẦU NGHIÊM NGẶT]:
        - CHỈ gõ lại nội dung có trong PDF
        - KHÔNG thêm bất kỳ nội dung nào khác
        - Giữ nguyên cấu trúc và định dạng của văn bản gốc, bỏ qua phần hình vẽ trong PDF
        - [Bắt buộc] tất cả công thức toán học viết dưới dạng LaTeX được bọc trong dấu $
        """

        self.parent_converter.start_conversion(prompt, result_widget=self.result_text,
                                             convert_button=self.convert_button,
                                             export_buttons=[self.export_word_button, self.export_pandoc_button])

    def export_to_word(self):
        if not hasattr(self.parent_converter, 'pdf_text') or not self.parent_converter.pdf_text:
            QMessageBox.warning(self, "Error", "Please convert PDF to text first.")
            return

        # Save dialog
        file_dialog = QFileDialog()
        output_path, _ = file_dialog.getSaveFileName(self, "Save Word Document", "", "Word Documents (*.docx)")

        if output_path:
            try:
                self.parent_converter.export_with_python_docx(output_path)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"An error occurred during export:\n{str(e)}")

    def export_to_word_pandoc(self):
        if not hasattr(self.parent_converter, 'pdf_text') or not self.parent_converter.pdf_text:
            QMessageBox.warning(self, "Error", "Please convert PDF to text first.")
            return

        # Save dialog
        file_dialog = QFileDialog()
        output_path, _ = file_dialog.getSaveFileName(self, "Save Word Document (Pandoc)", "", "Word Documents (*.docx)")

        if output_path:
            try:
                self.export_with_pandoc(output_path)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"An error occurred during pandoc export:\n{str(e)}")

    def export_with_pandoc(self, output_path):
        """Export using pandoc for better math formula handling"""
        try:
            # Create temporary markdown file
            temp_md = tempfile.mktemp(suffix='.md')

            # Write to temporary markdown file
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(self.parent_converter.pdf_text)

            # Run pandoc command
            cmd = [
                'pandoc',
                temp_md,
                '-o', output_path,
                '--from=markdown',
                '--to=docx',
                '--standalone'
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')

            if result.returncode == 0:
                QMessageBox.information(self, "Export Complete",
                                      f"Document exported successfully with Pandoc to:\n{output_path}\n\n"
                                      f"Math formulas should be properly rendered.")
            else:
                QMessageBox.warning(self, "Pandoc Error",
                                  f"Pandoc failed with error:\n{result.stderr}")

            # Clean up
            if os.path.exists(temp_md):
                os.unlink(temp_md)

        except FileNotFoundError:
            QMessageBox.warning(self, "Pandoc Not Found",
                              "Pandoc is not installed or not found in PATH.\n"
                              "Please install Pandoc from https://pandoc.org/")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"An error occurred during pandoc export:\n{str(e)}")

class ImageTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_converter = parent
        self.uploaded_images = []
        self.initUI()

    def initUI(self):
        layout = QVBoxLayout()

        # Title
        title_label = QLabel("Image to Word Converter")
        title_label.setStyleSheet("font-size: 16px; font-weight: bold; margin: 10px;")
        layout.addWidget(title_label)

        # Upload buttons
        upload_layout = QHBoxLayout()
        self.upload_images_button = QPushButton('Upload Images')
        self.upload_images_button.clicked.connect(self.upload_images)

        self.paste_from_clipboard_button = QPushButton('Paste from Clipboard')
        self.paste_from_clipboard_button.clicked.connect(self.paste_from_clipboard)

        self.clear_images_button = QPushButton('Clear All Images')
        self.clear_images_button.clicked.connect(self.clear_images)

        upload_layout.addWidget(self.upload_images_button)
        upload_layout.addWidget(self.paste_from_clipboard_button)
        upload_layout.addWidget(self.clear_images_button)
        layout.addLayout(upload_layout)

        # Image status
        self.image_status_label = QLabel("No images loaded")
        layout.addWidget(self.image_status_label)

        # Preview button
        self.preview_images_button = QPushButton('Preview Images')
        self.preview_images_button.clicked.connect(self.show_image_preview)
        self.preview_images_button.setEnabled(False)
        layout.addWidget(self.preview_images_button)

        # Convert button
        self.convert_button = QPushButton('Convert Images to Text')
        self.convert_button.clicked.connect(self.convert_images_to_text)
        self.convert_button.setEnabled(False)
        layout.addWidget(self.convert_button)

        # Export buttons
        export_layout = QHBoxLayout()
        self.export_word_button = QPushButton('Export to Word Document (python-docx)')
        self.export_word_button.clicked.connect(self.export_to_word)
        self.export_word_button.setEnabled(False)

        self.export_pandoc_button = QPushButton('Export to Word with Pandoc (Math formulas)')
        self.export_pandoc_button.clicked.connect(self.export_to_word_pandoc)
        self.export_pandoc_button.setEnabled(False)

        export_layout.addWidget(self.export_word_button)
        export_layout.addWidget(self.export_pandoc_button)
        layout.addLayout(export_layout)

        # Result display
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        layout.addWidget(QLabel('Results:'))
        layout.addWidget(self.result_text)

        self.setLayout(layout)

    def upload_images(self):
        file_dialog = QFileDialog()
        file_paths, _ = file_dialog.getOpenFileNames(
            self,
            "Select Image files",
            "",
            "Image Files (*.png *.jpg *.jpeg *.bmp *.gif *.tiff)"
        )

        if file_paths:
            self.add_images_from_files(file_paths)

    def paste_from_clipboard(self):
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            image = clipboard.image()
            if not image.isNull():
                # Save clipboard image to temp file
                temp_path = tempfile.mktemp(suffix='.png')
                if image.save(temp_path, 'PNG'):
                    self.add_images_from_files([temp_path], source='Clipboard')
                    QMessageBox.information(self, "Success", "Image pasted from clipboard!")
                else:
                    QMessageBox.warning(self, "Error", "Failed to save clipboard image.")
            else:
                QMessageBox.warning(self, "Error", "No valid image in clipboard.")
        else:
            QMessageBox.warning(self, "Error", "No image found in clipboard.")

    def add_images_from_files(self, file_paths, source='Upload'):
        for file_path in file_paths:
            try:
                # Create a copy in the output directory if needed
                if self.parent_converter.output_dir:
                    img_dir = os.path.join(self.parent_converter.output_dir, "Images")
                    if not os.path.exists(img_dir):
                        os.makedirs(img_dir)

                    filename = f"img_{len(self.uploaded_images) + 1}_{os.path.basename(file_path)}"
                    dest_path = os.path.join(img_dir, filename)
                    shutil.copy2(file_path, dest_path)

                    # Get image size info
                    try:
                        with Image.open(dest_path) as pil_img:
                            width, height = pil_img.size
                            size_info = f"{width}x{height}px"
                    except:
                        size_info = "Unknown"

                    self.uploaded_images.append({
                        'path': dest_path,
                        'filename': filename,
                        'source': source,
                        'size_info': size_info,
                        'index': len(self.uploaded_images) + 1
                    })
                else:
                    # If no output directory, use original path
                    try:
                        with Image.open(file_path) as pil_img:
                            width, height = pil_img.size
                            size_info = f"{width}x{height}px"
                    except:
                        size_info = "Unknown"

                    self.uploaded_images.append({
                        'path': file_path,
                        'filename': os.path.basename(file_path),
                        'source': source,
                        'size_info': size_info,
                        'index': len(self.uploaded_images) + 1
                    })

            except Exception as e:
                QMessageBox.warning(self, "Error", f"Failed to process {file_path}:\n{str(e)}")

        self.update_image_status()

    def clear_images(self):
        self.uploaded_images = []
        self.update_image_status()

    def update_image_status(self):
        if self.uploaded_images:
            self.image_status_label.setText(f"Loaded {len(self.uploaded_images)} images")
            self.preview_images_button.setEnabled(True)
            self.convert_button.setEnabled(True)
        else:
            self.image_status_label.setText("No images loaded")
            self.preview_images_button.setEnabled(False)
            self.convert_button.setEnabled(False)
            self.export_word_button.setEnabled(False)
            self.export_pandoc_button.setEnabled(False)

    def show_image_preview(self):
        if not self.uploaded_images:
            QMessageBox.information(self, "Info", "No images to preview.")
            return

        dialog = ImagePreviewDialog(self.uploaded_images, self)
        dialog.exec()

    def convert_images_to_text(self):
        if not self.parent_converter.client:
            QMessageBox.warning(self, "Error", "Please set the API Key first.")
            return

        if not self.uploaded_images:
            self.result_text.setText("Please upload images first.")
            return

        # Upload all images to Gemini
        try:
            uploaded_files = []
            for img_info in self.uploaded_images:
                uploaded_file = self.parent_converter.client.files.upload(file=img_info['path'])
                uploaded_files.append(uploaded_file)

            prompt = f"""
            Hãy nhận diện và gõ lại [CHÍNH XÁC] toàn bộ nội dung trong {len(self.uploaded_images)} hình ảnh thành văn bản, bao gồm tất cả công thức Toán học được bọc trong dấu $.

            [QUY TẮC NGHIÊM NGẶT]:
            - CHỈ gõ lại nội dung có trong hình ảnh
            - KHÔNG thêm bất kỳ nội dung nào khác
            - Giữ nguyên cấu trúc và định dạng của văn bản gốc, bỏ qua phần hình vẽ trong PDF
            - [Bắt buộc] tất cả công thức toán học viết dưới dạng LaTeX được bọc trong dấu $

            """

            # Create content list with all uploaded files
            content_list = uploaded_files + [prompt]

            # Start conversion with multiple images
            self.start_image_conversion(content_list)

        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to upload images: {str(e)}")

    def start_image_conversion(self, content_list):
        self.convert_button.setEnabled(False)
        self.export_word_button.setEnabled(False)
        self.export_pandoc_button.setEnabled(False)

        self.parent_converter.progress_bar.setValue(0)
        self.result_text.clear()
        self.result_text.append("Starting image conversion process...")
        self.parent_converter.status_label.setText("Status: Converting images...")

        # Create a special thread for image conversion
        self.image_conversion_thread = ImageConversionThread(
            self.parent_converter.client,
            content_list
        )
        self.image_conversion_thread.progress.connect(self.parent_converter.update_progress)
        self.image_conversion_thread.finished.connect(self.on_conversion_finished)
        self.image_conversion_thread.error.connect(self.on_conversion_error)
        self.image_conversion_thread.start()

    def on_conversion_finished(self, text):
        # Save processed text
        self.parent_converter.pdf_text = self.parent_converter.process_formulas(text)

        self.result_text.clear()
        self.result_text.append("Images converted successfully. Here's the content:\n\n")
        self.result_text.append(self.parent_converter.pdf_text)

        self.convert_button.setEnabled(True)
        self.export_word_button.setEnabled(True)
        self.export_pandoc_button.setEnabled(True)

        self.parent_converter.status_label.setText("Status: Image conversion completed")

    def on_conversion_error(self, error_message):
        self.result_text.append(f"An error occurred during conversion: {error_message}")
        self.convert_button.setEnabled(True)
        self.parent_converter.status_label.setText("Status: Error occurred")

    def export_to_word(self):
        if not hasattr(self.parent_converter, 'pdf_text') or not self.parent_converter.pdf_text:
            QMessageBox.warning(self, "Error", "Please convert images to text first.")
            return

        # Save dialog
        file_dialog = QFileDialog()
        output_path, _ = file_dialog.getSaveFileName(self, "Save Word Document", "", "Word Documents (*.docx)")

        if output_path:
            try:
                self.parent_converter.export_with_python_docx(output_path, include_original_images=True)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"An error occurred during export:\n{str(e)}")

    def export_to_word_pandoc(self):
        if not hasattr(self.parent_converter, 'pdf_text') or not self.parent_converter.pdf_text:
            QMessageBox.warning(self, "Error", "Please convert images to text first.")
            return

        # Save dialog
        file_dialog = QFileDialog()
        output_path, _ = file_dialog.getSaveFileName(self, "Save Word Document (Pandoc)", "", "Word Documents (*.docx)")

        if output_path:
            try:
                self.export_with_pandoc(output_path)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"An error occurred during pandoc export:\n{str(e)}")

    def export_with_pandoc(self, output_path):
        """Export using pandoc for better math formula handling"""
        try:
            # Create temporary markdown file
            temp_md = tempfile.mktemp(suffix='.md')

            # Write to temporary markdown file
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(self.parent_converter.pdf_text)

            # Run pandoc command
            cmd = [
                'pandoc',
                temp_md,
                '-o', output_path,
                '--from=markdown',
                '--to=docx',
                '--standalone'
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')

            if result.returncode == 0:
                QMessageBox.information(self, "Export Complete",
                                      f"Document exported successfully with Pandoc to:\n{output_path}\n\n"
                                      f"Math formulas should be properly rendered.")
            else:
                QMessageBox.warning(self, "Pandoc Error",
                                  f"Pandoc failed with error:\n{result.stderr}")

            # Clean up
            if os.path.exists(temp_md):
                os.unlink(temp_md)

        except FileNotFoundError:
            QMessageBox.warning(self, "Pandoc Not Found",
                              "Pandoc is not installed or not found in PATH.\n"
                              "Please install Pandoc from https://pandoc.org/")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"An error occurred during pandoc export:\n{str(e)}")

class ImageConversionThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, client, content_list):
        super().__init__()
        self.client = client
        self.content_list = content_list
        self.max_retries = 3
        self.retry_delay = 60
        self.is_running = True

    def run(self):
        for attempt in range(self.max_retries):
            if not self.is_running:
                return
            try:
                # Simulate progress
                for i in range(0, 51):
                    if not self.is_running:
                        return
                    self.progress.emit(i)
                    self.msleep(50)

                # Progress simulation during generation
                for i in range(51, 101):
                    if not self.is_running:
                        return
                    self.progress.emit(i)
                    self.msleep(30)

                # Generate content using new simplified API
                response = self.client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=self.content_list,
                )

                self.finished.emit(response.text)
                return

            except Exception as e:
                if "429" in str(e) and attempt < self.max_retries - 1:
                    self.error.emit(f"Rate limit exceeded. Retrying in {self.retry_delay} seconds...")
                    self.msleep(self.retry_delay * 1000)
                else:
                    self.error.emit(str(e))
                    return

    def stop(self):
        self.is_running = False

class PDFToTextConverter(QWidget):
    def __init__(self):
        super().__init__()
        self.api_key = ""
        self.file_path = None
        self.uploaded_file = None
        self.client = None
        self.pdf_text = ""
        self.output_dir = ""
        self.initUI()
        self.load_api_key()
        self.conversion_thread = None

    def initUI(self):
        main_layout = QVBoxLayout()

        # API Key section
        api_layout = QHBoxLayout()
        self.api_key_input = QLineEdit()
        self.api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.api_key_button = QPushButton('Set API Key')
        self.api_key_button.clicked.connect(self.set_api_key)
        self.edit_api_key_button = QPushButton('Edit API Key')
        self.edit_api_key_button.clicked.connect(self.edit_api_key)
        self.edit_api_key_button.setEnabled(False)
        api_layout.addWidget(QLabel('Gemini API Key:'))
        api_layout.addWidget(self.api_key_input)
        api_layout.addWidget(self.api_key_button)
        api_layout.addWidget(self.edit_api_key_button)
        main_layout.addLayout(api_layout)

        # File upload section
        upload_layout = QHBoxLayout()
        self.upload_pdf_button = QPushButton('Upload PDF')
        self.upload_pdf_button.clicked.connect(self.upload_pdf)
        self.file_label = QLabel('No PDF file selected')
        upload_layout.addWidget(self.upload_pdf_button)
        upload_layout.addWidget(self.file_label)
        main_layout.addLayout(upload_layout)

        # Progress bar
        self.progress_bar = QProgressBar(self)
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        main_layout.addWidget(QLabel("Progress:"))
        main_layout.addWidget(self.progress_bar)

        # Status label
        self.status_label = QLabel("Status: Idle")
        main_layout.addWidget(self.status_label)

        # Create tab widget
        self.tab_widget = QTabWidget()

        # Create tabs
        self.word_tab = WordTab(self)
        self.image_tab = ImageTab(self)

        # Add tabs
        self.tab_widget.addTab(self.word_tab, "PDF to Word")
        self.tab_widget.addTab(self.image_tab, "Image to Word")

        main_layout.addWidget(self.tab_widget)

        self.setLayout(main_layout)
        self.setWindowTitle('PDF & Image to Word Converter')
        self.setGeometry(300, 300, 1200, 900)

    def set_api_key(self):
        self.api_key = self.api_key_input.text()
        if self.api_key:
            self.setup_client()
            self.api_key_input.setEnabled(False)
            self.api_key_button.setEnabled(False)
            self.edit_api_key_button.setEnabled(True)
            self.save_api_key()
            QMessageBox.information(self, "Success", "API Key set successfully!")
        else:
            QMessageBox.warning(self, "Error", "Please enter an API Key.")

    def setup_client(self):
        try:
            self.client = genai.Client(api_key=self.api_key)
            print("Gemini client initialized successfully")
        except Exception as e:
            print(f"Error setting up client: {e}")
            self.client = None

    def edit_api_key(self):
        self.api_key_input.setEnabled(True)
        self.api_key_button.setEnabled(True)
        self.edit_api_key_button.setEnabled(False)

    def save_api_key(self):
        try:
            with open('api_key.txt', 'w') as f:
                f.write(self.api_key)
        except Exception as e:
            print(f"Error saving API key: {str(e)}")

    def load_api_key(self):
        try:
            if os.path.exists('api_key.txt'):
                with open('api_key.txt', 'r') as f:
                    self.api_key = f.read().strip()
                    self.api_key_input.setText(self.api_key)
                    self.set_api_key()
        except Exception as e:
            print(f"Error loading API key: {str(e)}")

    def upload_pdf(self):
        if not self.client:
            QMessageBox.warning(self, "Error", "Please set the API Key first.")
            return

        file_dialog = QFileDialog()
        self.file_path, _ = file_dialog.getOpenFileName(self, "Select PDF file", "", "PDF Files (*.pdf)")

        if self.file_path:
            file_name = os.path.basename(self.file_path)
            self.file_label.setText(f"File: {file_name}")

            # Set output directory to the same directory as the PDF file
            self.output_dir = os.path.dirname(self.file_path)

            self.process_pdf()

    def process_pdf(self):
        self.status_label.setText("Status: Processing PDF...")

        try:
            # Upload PDF using new API
            safe_path = self.file_path.replace("\\", "/")

            # Upload file using new API
            self.uploaded_file = self.client.files.upload(file=safe_path)

            print(f"File uploaded successfully: {self.uploaded_file.uri}")

            # Enable convert button in PDF tab
            self.word_tab.convert_button.setEnabled(True)

            self.status_label.setText("Status: PDF ready for conversion")

        except Exception as e:
            print(f"Error processing PDF: {e}")
            QMessageBox.warning(self, "Error", f"Failed to upload PDF: {str(e)}")

    def start_conversion(self, prompt, result_widget=None, convert_button=None, export_buttons=None):
        if convert_button:
            convert_button.setEnabled(False)
        if export_buttons:
            for btn in export_buttons:
                btn.setEnabled(False)

        self.progress_bar.setValue(0)
        if result_widget:
            result_widget.clear()
            result_widget.append("Starting conversion process...")
        self.status_label.setText("Status: Converting...")

        # Start conversion thread
        self.conversion_thread = ConversionThread(self.client, self.uploaded_file, prompt)
        self.conversion_thread.progress.connect(self.update_progress)
        self.conversion_thread.finished.connect(
            lambda text: self.on_conversion_finished(text, result_widget, convert_button, export_buttons)
        )
        self.conversion_thread.error.connect(
            lambda error: self.on_conversion_error(error, result_widget, convert_button)
        )
        self.conversion_thread.start()

    def update_progress(self, value):
        self.progress_bar.setValue(value)

    def on_conversion_finished(self, text, result_widget=None, convert_button=None, export_buttons=None):
        # Save processed text
        self.pdf_text = self.process_formulas(text)

        if result_widget:
            result_widget.clear()
            result_widget.append("PDF converted successfully. Here's the content:\n\n")
            result_widget.append(self.pdf_text)

        if convert_button:
            convert_button.setEnabled(True)
        if export_buttons:
            for btn in export_buttons:
                btn.setEnabled(True)

        self.status_label.setText("Status: Conversion completed")

    def on_conversion_error(self, error_message, result_widget=None, convert_button=None):
        if result_widget:
            result_widget.append(f"An error occurred during conversion: {error_message}")
        if convert_button:
            convert_button.setEnabled(True)
        self.status_label.setText("Status: Error occurred")

    def process_formulas(self, text):
        def process_math_content(match):
            content = match.group(1)
            content = content.replace('π', '\\pi')
            content = re.sub(r'√(\d+)', r'\\sqrt{\1}', content)
            content = re.sub(r'√\{([^}]+)\}', r'\\sqrt{\1}', content)
            content = content.replace('≠', '\\neq')
            content = content.replace('*', '')
            return f'${content}$'

        text = re.sub(r'\$(.+?)\$', process_math_content, text, flags=re.DOTALL)
        return text

    def export_with_python_docx(self, output_path, include_original_images=False):
        """Export using python-docx with optional original images"""
        doc = Document()
        doc.add_heading('Converted Document', 0)

        lines = self.pdf_text.split('\n')
        i = 0

        # If this is from image tab and we want to include original images
        if include_original_images and hasattr(self.image_tab, 'uploaded_images'):
            doc.add_heading('Original Images', level=1)
            for img_info in self.image_tab.uploaded_images:
                try:
                    if os.path.exists(img_info['path']):
                        doc.add_paragraph(f"Image {img_info['index']}: {img_info['filename']}")
                        doc.add_picture(img_info['path'], width=Inches(5))
                        doc.add_paragraph("")  # spacing
                except Exception as e:
                    print(f"Error adding original image {img_info['index']}: {e}")
                    doc.add_paragraph(f"[Error: Could not insert image {img_info['index']}]")

            doc.add_heading('Extracted Text', level=1)

        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # Check for Markdown table
            if '|' in line and len(line.split('|')) > 2:
                table_lines = []
                j = i

                # Collect consecutive table lines
                while j < len(lines):
                    current_line = lines[j].strip()
                    if '|' in current_line and len(current_line.split('|')) > 2:
                        # Skip separator lines like |---|---|
                        separator_pattern = r'^\s*\|[\s\-\:]*\|\s*$'
                        if not re.match(separator_pattern, current_line):
                            table_lines.append(current_line)
                        j += 1
                    else:
                        break

                if table_lines:
                    print(f"Creating table with {len(table_lines)} rows")
                    self.create_word_table(doc, table_lines)
                    i = j
                else:
                    doc.add_paragraph(line)
                    i += 1
            else:
                # Regular text
                doc.add_paragraph(line)
                i += 1

        doc.save(output_path)

        success_msg = (f"Document exported successfully to:\n{output_path}\n\n"
                      f"Statistics:\n"
                      f"- Processed {len(lines)} lines of text")

        if include_original_images and hasattr(self.image_tab, 'uploaded_images'):
            success_msg += f"\n- Included {len(self.image_tab.uploaded_images)} original images"

        print(success_msg)
        QMessageBox.information(self, "Export Complete", success_msg)

    def create_word_table(self, doc, table_lines):
        """Create a Word table from Markdown table lines"""
        if not table_lines:
            return

        print(f"Creating table from {len(table_lines)} lines:")
        for i, line in enumerate(table_lines):
            print(f"  Line {i}: {repr(line)}")

        # Parse table data
        table_data = []
        for line_num, line in enumerate(table_lines):
            if not line.strip():
                continue

            # Split by | and clean up cells
            raw_cells = line.split('|')
            cells = []

            for cell in raw_cells:
                cleaned_cell = cell.strip()
                if cleaned_cell or (cells and line_num < len(table_lines) - 1):
                    cells.append(cleaned_cell)

            # Remove leading/trailing empty cells only
            while cells and not cells[0]:
                cells.pop(0)
            while cells and not cells[-1]:
                cells.pop()

            if cells:
                table_data.append(cells)
                print(f"  Parsed row {len(table_data)}: {cells}")

        if not table_data:
            print("No valid table data found, adding as regular text")
            for line in table_lines:
                doc.add_paragraph(line)
            return

        try:
            rows = len(table_data)
            cols = max(len(row) for row in table_data) if table_data else 0

            print(f"Creating {rows}x{cols} table")

            if rows > 0 and cols > 0:
                # Create table
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'

                # Fill table with data
                for i, row_data in enumerate(table_data):
                    if i >= len(table.rows):
                        print(f"Warning: Row {i} exceeds table rows")
                        break

                    row = table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        if j >= len(row.cells):
                            print(f"Warning: Column {j} exceeds table columns in row {i}")
                            break

                        try:
                            clean_data = str(cell_data).replace('\n', ' ').replace('\r', '')
                            row.cells[j].text = clean_data
                            print(f"  Cell [{i}][{j}]: {repr(clean_data)}")
                        except Exception as cell_error:
                            print(f"Error setting cell [{i}][{j}]: {cell_error}")
                            row.cells[j].text = str(cell_data)

                doc.add_paragraph("")  # spacing after table
                print("Table created successfully")

            else:
                print(f"Invalid table dimensions: {rows}x{cols}")
                for line in table_lines:
                    doc.add_paragraph(line)

        except Exception as e:
            print(f"Error creating table: {e}")
            print(f"Table data that caused error: {table_data}")
            doc.add_paragraph("Table conversion failed, showing as text:")
            for line in table_lines:
                doc.add_paragraph(f"  {line}")

    def closeEvent(self, event):
        self.cleanup_and_close()
        super().closeEvent(event)

    def cleanup_and_close(self):
        # Stop conversion thread if running
        if self.conversion_thread and self.conversion_thread.isRunning():
            self.conversion_thread.stop()
            self.conversion_thread.wait()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = PDFToTextConverter()
    ex.show()
    sys.exit(app.exec())